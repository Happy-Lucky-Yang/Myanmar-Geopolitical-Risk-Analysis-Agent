"""
analyzer.report_generator - 自动化结构化报告生成

申报书要求: "自动化生成结构化报告（docx/html），包含量化指标与趋势研判"

功能:
  - 基于分析结果自动生成 HTML 报告 (Jinja2 模板)
  - 可选 docx 输出 (python-docx 库)
  - 报告结构: 标题 + 时间范围 + 风险评分摘要 + 关键事件列表 + 趋势图表 + 预测研判

API: /api/report?format=html|docx&days=30
"""
import os
import io
import logging
import threading
from datetime import datetime, timedelta
from typing import Dict, Optional

logger = logging.getLogger(__name__)

# Jinja2 HTML 报告模板 (内嵌, 避免额外文件依赖)
_HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<title>{{ title }}</title>
<style>
  * { margin: 0; padding: 0; box-sizing: border-box; }
  body { font-family: "Microsoft YaHei", "Segoe UI", sans-serif;
         background: #0d1117; color: #c9d1d9; padding: 40px; line-height: 1.6; }
  .report { max-width: 900px; margin: 0 auto; }
  h1 { color: #58a6ff; border-bottom: 2px solid #30363d; padding-bottom: 12px; margin-bottom: 24px; font-size: 24px; }
  h2 { color: #8b949e; margin: 28px 0 12px; font-size: 18px; border-left: 3px solid #58a6ff; padding-left: 10px; }
  .meta { color: #8b949e; font-size: 13px; margin-bottom: 20px; }
  .card { background: #161b22; border: 1px solid #30363d; border-radius: 8px;
          padding: 20px; margin-bottom: 16px; }
  .score-big { font-size: 48px; font-weight: bold; text-align: center; }
  .score-high { color: #f85149; }
  .score-mid { color: #d29922; }
  .score-low { color: #3fb950; }
  .metric-grid { display: grid; grid-template-columns: repeat(auto-fit, minmax(180px, 1fr));
                 gap: 12px; margin: 16px 0; }
  .metric-item { background: #0d1117; border: 1px solid #30363d; border-radius: 6px;
                 padding: 12px; text-align: center; }
  .metric-value { font-size: 24px; font-weight: bold; color: #58a6ff; }
  .metric-label { font-size: 12px; color: #8b949e; margin-top: 4px; }
  table { width: 100%; border-collapse: collapse; margin: 12px 0; }
  th, td { padding: 10px 12px; text-align: left; border-bottom: 1px solid #21262d; }
  th { color: #8b949e; font-size: 12px; text-transform: uppercase; }
  .trend-up { color: #f85149; }
  .trend-down { color: #3fb950; }
  .trend-flat { color: #8b949e; }
  .footer { margin-top: 40px; padding-top: 20px; border-top: 1px solid #30363d;
            color: #484f58; font-size: 11px; text-align: center; }
  .disclaimer { background: #1c1208; border: 1px solid #d29922; border-radius: 6px;
                padding: 12px; margin: 16px 0; font-size: 12px; color: #d29922; }
  @media print { body { background: #fff; color: #1c1e21; }
    .card { border-color: #d0d7de; } h1 { color: #0969da; } }
</style>
</head>
<body>
<div class="report">
  <h1>{{ title }}</h1>
  <div class="meta">
    报告时间范围: {{ date_range }} | 生成时间: {{ generated_at }} | 数据来源: {{ sources }}
  </div>

  {% if risk_summary %}
  <h2>风险评分摘要</h2>
  <div class="card">
    <div class="score-big {{ risk_summary.score_class }}">
      {{ risk_summary.score }}
    </div>
    <p style="text-align:center; color:#8b949e; margin-top:8px;">
      {{ risk_summary.level }} | {{ risk_summary.trend_text }}
    </p>
    <div class="metric-grid">
      {% for m in risk_summary.metrics %}
      <div class="metric-item">
        <div class="metric-value">{{ m.value }}</div>
        <div class="metric-label">{{ m.name }}</div>
      </div>
      {% endfor %}
    </div>
  </div>
  {% endif %}

  {% if key_events %}
  <h2>关键事件</h2>
  <div class="card">
    <table>
      <thead><tr><th>日期</th><th>事件</th><th>类型</th><th>严重程度</th></tr></thead>
      <tbody>
      {% for ev in key_events %}
      <tr>
        <td>{{ ev.date }}</td>
        <td>{{ ev.title }}</td>
        <td>{{ ev.event_type }}</td>
        <td>{{ ev.severity }}</td>
      </tr>
      {% endfor %}
      </tbody>
    </table>
  </div>
  {% endif %}

  {% if trend_data %}
  <h2>趋势分析</h2>
  <div class="card">
    <p>分析周期: {{ trend_data.period }}</p>
    <p>趋势方向: <span class="{{ trend_data.trend_class }}">{{ trend_data.trend }}</span></p>
    <p>数据点数: {{ trend_data.data_points }}</p>
    {% if trend_data.forecast %}
    <p>7日预测: {{ trend_data.forecast_summary }}</p>
    <p>预测置信度: {{ trend_data.confidence }}</p>
    {% endif %}
  </div>
  {% endif %}

  {% if economic_data %}
  <h2>宏观经济指标</h2>
  <div class="card">
    <div class="metric-grid">
      {% for ind in economic_data %}
      <div class="metric-item">
        <div class="metric-value">{{ ind.value }}</div>
        <div class="metric-label">{{ ind.name }}</div>
      </div>
      {% endfor %}
    </div>
    <div class="disclaimer">经济指标数据来源于 World Bank Open Data，标注"{{ economic_quality }}"的数据为API直接获取，标注"估算"的为模型推断值。</div>
  </div>
  {% endif %}

  {% if assessment %}
  <h2>综合研判与建议</h2>
  <div class="card">
    <p>{{ assessment }}</p>
  </div>
  {% endif %}

  <div class="footer">
    华东师范大学 地缘环境智能计算实验室 | 本报告由系统自动生成，仅供参考
  </div>
</div>
</body>
</html>"""


class ReportGenerator:
    """自动化结构化报告生成器"""

    def __init__(self):
        self._jinja_env = None

    def _get_jinja_env(self):
        """延迟初始化 Jinja2 环境"""
        if self._jinja_env is None:
            try:
                from jinja2 import Environment
                self._jinja_env = Environment()
            except ImportError:
                logger.error("[Report] Jinja2 未安装")
                raise
        return self._jinja_env

    def generate_html_report(self, days: int = 30) -> str:
        """
        生成 HTML 格式结构化报告

        :param days: 分析最近多少天的数据
        :return: HTML 字符串
        """
        from analyzer.data_loader import get_data_loader
        from analyzer.trend import get_trend_analyzer

        loader = get_data_loader()
        history = loader.load_risk_history(days=days)

        # 收集报告数据
        context = self._build_report_context(days, history)

        # 渲染模板
        env = self._get_jinja_env()
        template = env.from_string(_HTML_TEMPLATE)
        return template.render(**context)

    def generate_docx_report(self, days: int = 30) -> bytes:
        """
        生成 DOCX 格式报告

        :param days: 分析最近多少天的数据
        :return: docx 字节流
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise RuntimeError("python-docx 未安装，请运行: pip install python-docx")

        from analyzer.data_loader import get_data_loader
        loader = get_data_loader()
        history = loader.load_risk_history(days=days)

        context = self._build_report_context(days, history)

        doc = Document()

        # 标题
        title = doc.add_heading(context["title"], level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 元信息
        doc.add_paragraph(
            f"报告时间范围: {context['date_range']} | "
            f"生成时间: {context['generated_at']}"
        )

        # 风险评分摘要
        risk = context.get("risk_summary")
        if risk:
            doc.add_heading("风险评分摘要", level=2)
            p = doc.add_paragraph()
            run = p.add_run(f"综合风险分: {risk['score']} ({risk['level']})")
            run.bold = True
            doc.add_paragraph(f"趋势: {risk['trend_text']}")

            # 指标表
            table = doc.add_table(rows=1, cols=2)
            table.style = "Light Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "指标"
            hdr[1].text = "值"
            for m in risk.get("metrics", []):
                row = table.add_row().cells
                row[0].text = m["name"]
                row[1].text = str(m["value"])

        # 关键事件
        events = context.get("key_events", [])
        if events:
            doc.add_heading("关键事件", level=2)
            table = doc.add_table(rows=1, cols=4)
            table.style = "Light Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "日期"
            hdr[1].text = "事件"
            hdr[2].text = "类型"
            hdr[3].text = "严重程度"
            for ev in events[:10]:
                row = table.add_row().cells
                row[0].text = ev["date"]
                row[1].text = ev["title"]
                row[2].text = ev["event_type"]
                row[3].text = ev["severity"]

        # 趋势分析
        trend = context.get("trend_data")
        if trend:
            doc.add_heading("趋势分析", level=2)
            doc.add_paragraph(f"分析周期: {trend['period']}")
            doc.add_paragraph(f"趋势方向: {trend['trend']}")
            if trend.get("forecast"):
                doc.add_paragraph(f"7日预测: {trend['forecast_summary']}")
                doc.add_paragraph(f"预测置信度: {trend['confidence']}")

        # 经济数据
        econ = context.get("economic_data")
        if econ:
            doc.add_heading("宏观经济指标", level=2)
            for ind in econ:
                doc.add_paragraph(f"• {ind['name']}: {ind['value']}")

        # 综合研判
        assessment = context.get("assessment")
        if assessment:
            doc.add_heading("综合研判与建议", level=2)
            doc.add_paragraph(assessment)

        # 页脚
        doc.add_paragraph()
        footer_p = doc.add_paragraph(
            "华东师范大学 地缘环境智能计算实验室 | 本报告由系统自动生成"
        )
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 保存到字节流
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    # ============================================================
    # 内部: 构建报告上下文数据
    # ============================================================

    def _build_report_context(self, days: int, history: list) -> Dict:
        """构建报告模板渲染所需的上下文数据"""
        now = datetime.now()
        start_date = (now - timedelta(days=days)).strftime("%Y-%m-%d")
        end_date = now.strftime("%Y-%m-%d")

        context = {
            "title": f"缅甸地缘风险分析报告 ({start_date} ~ {end_date})",
            "date_range": f"{start_date} 至 {end_date}",
            "generated_at": now.strftime("%Y-%m-%d %H:%M"),
            "sources": "新闻文本 + GDELT + World Bank + VIIRS 遥感",
        }

        # 风险评分摘要
        if history:
            scores = [r["risk_score"] for r in history]
            avg_score = sum(scores) / len(scores)
            latest_score = scores[-1]

            score_class = "score-high" if latest_score >= 70 else "score-mid" if latest_score >= 40 else "score-low"
            level = "高风险" if latest_score >= 70 else "中风险" if latest_score >= 40 else "低风险"

            # 趋势判断
            if len(scores) >= 7:
                recent_avg = sum(scores[-7:]) / 7
                older_avg = sum(scores[:7]) / min(7, len(scores[:7]))
                if recent_avg > older_avg + 5:
                    trend_text = "近期风险呈上升趋势"
                elif recent_avg < older_avg - 5:
                    trend_text = "近期风险呈下降趋势"
                else:
                    trend_text = "近期风险相对平稳"
            else:
                trend_text = "数据不足，无法判断趋势"

            # 收集指标
            last_details = history[-1].get("details", {})
            metrics = []
            metric_names = {
                "conflict_frequency": "冲突频次",
                "sentiment_avg": "情感风险",
                "nightlight_change": "夜光变化",
                "refugee_change": "难民变化",
                "event_severity": "事件严重度",
            }
            for key, label in metric_names.items():
                val = last_details.get(key)
                if val is not None:
                    metrics.append({"name": label, "value": f"{val:.2f}"})

            context["risk_summary"] = {
                "score": f"{latest_score:.1f}",
                "score_class": score_class,
                "level": level,
                "trend_text": trend_text,
                "metrics": metrics,
            }

            # 关键事件 (从分析结果中提取)
            context["key_events"] = self._extract_key_events(history)
        else:
            context["risk_summary"] = None
            context["key_events"] = []

        # 趋势分析
        context["trend_data"] = self._build_trend_data(history)

        # 经济数据
        context["economic_data"], context["economic_quality"] = self._build_economic_data()

        # 综合研判
        context["assessment"] = self._generate_assessment(context)

        return context

    def _extract_key_events(self, history: list) -> list:
        """从历史风险记录中提取关键事件 (高风险日)"""
        events = []
        for record in history:
            score = record.get("risk_score", 0)
            if score >= 60:
                events.append({
                    "date": record.get("date", ""),
                    "title": f"风险评分 {score:.1f}",
                    "event_type": "高风险预警",
                    "severity": "高" if score >= 80 else "中",
                })
        # 只返回最近 10 个
        return events[-10:]

    def _build_trend_data(self, history: list) -> Optional[Dict]:
        """构建趋势分析数据"""
        if not history or len(history) < 3:
            return None

        try:
            from analyzer.trend import get_trend_analyzer
            scores = [r["risk_score"] for r in history]
            trend_analyzer = get_trend_analyzer()
            trend_result = trend_analyzer.full_analysis(scores)
            forecast_result = trend_analyzer.forecast(scores, days_ahead=7)

            trend = trend_result.get("trend", "未知")
            trend_class = "trend-up" if trend == "上升" else "trend-down" if trend == "下降" else "trend-flat"

            forecast = forecast_result.get("forecast", [])
            if forecast:
                avg_forecast = sum(forecast) / len(forecast)
                forecast_summary = f"未来7日均值预计 {avg_forecast:.1f}"
            else:
                forecast_summary = "数据不足"

            return {
                "period": f"{history[0].get('date', '?')} ~ {history[-1].get('date', '?')}",
                "trend": trend,
                "trend_class": trend_class,
                "data_points": len(history),
                "forecast": forecast,
                "forecast_summary": forecast_summary,
                "confidence": f"{forecast_result.get('confidence', 0):.0f}%",
            }
        except Exception as e:
            logger.warning(f"[Report] 趋势数据构建失败: {e}")
            return None

    def _build_economic_data(self):
        """构建经济指标数据"""
        try:
            from data.economic_crawler import get_economic_crawler
            econ = get_economic_crawler()
            data = econ.get_all_indicators()

            indicators = [
                {"name": "GDP 增长率", "value": f"{data.get('gdp_growth', 0):.1f}%"},
                {"name": "人均 GDP", "value": f"${data.get('gdp_per_capita', 0):,.0f}"},
                {"name": "通胀率", "value": f"{data.get('inflation', 0):.1f}%"},
                {"name": "贸易占 GDP", "value": f"{data.get('trade_pct_gdp', 0):.1f}%"},
                {"name": "难民变化率", "value": f"{data.get('refugee_change', 0):.2f}"},
            ]
            quality = data.get("data_quality", "估算")
            return indicators, quality
        except Exception as e:
            logger.warning(f"[Report] 经济数据获取失败: {e}")
            return None, None

    def _generate_assessment(self, context: Dict) -> str:
        """基于数据生成简要综合研判"""
        parts = []

        risk = context.get("risk_summary")
        if risk:
            parts.append(f"当前缅甸地缘风险综合评分为 {risk['score']} 分，"
                         f"处于{risk['level']}水平。{risk['trend_text']}。")

        trend = context.get("trend_data")
        if trend:
            parts.append(f"过去 {trend['data_points']} 天的趋势分析显示风险{trend['trend']}，"
                         f"{trend.get('forecast_summary', '')}。")

        econ = context.get("economic_data")
        if econ:
            parts.append(f"经济方面，{econ[0]['name']}为{econ[0]['value']}，"
                         f"{econ[2]['name']}为{econ[2]['value']}。")

        if not parts:
            return "数据不足，暂无法生成综合研判。建议等待系统采集更多数据后再生成报告。"

        return " ".join(parts)


# ============================================================
# 单例
# ============================================================
_instance = None
_lock = threading.Lock()


def get_report_generator() -> ReportGenerator:
    """获取全局报告生成器单例"""
    global _instance
    if _instance is None:
        with _lock:
            if _instance is None:
                _instance = ReportGenerator()
    return _instance
